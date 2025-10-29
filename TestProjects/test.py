import os
import io
import uuid
import shutil
import logging
from datetime import datetime
from flask import (
    Flask,
    request,
    jsonify,
    render_template_string,
    send_from_directory,
    redirect,
    url_for,
    send_file,
)
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader, PdfWriter
import fitz
from PIL import Image
from docx import Document
import pdfplumber
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
import pikepdf


app = Flask(__name__)

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
app.config["UPLOAD_FOLDER"] = os.path.join(BASE_DIR, "uploads")
app.config["SAVED_FOLDER"] = os.path.join(BASE_DIR, "saved")
app.config["THUMBNAIL_FOLDER"] = os.path.join(BASE_DIR, "thumbnails")
app.config["RESULT_FOLDER"] = os.path.join(BASE_DIR, "comparison_results")

SPLIT_FOLDER = os.path.join(BASE_DIR, "split")
MERGE_FOLDER = os.path.join(BASE_DIR, "merge")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "output")

app.config["OUTPUT_FOLDER"] = OUTPUT_FOLDER

FOLDERS_TO_CLEAR = [
    app.config["UPLOAD_FOLDER"],
    app.config["SAVED_FOLDER"],
    app.config["THUMBNAIL_FOLDER"],
    app.config["RESULT_FOLDER"],
    SPLIT_FOLDER,
    MERGE_FOLDER,
    OUTPUT_FOLDER,
]

for folder in FOLDERS_TO_CLEAR:
    os.makedirs(folder, exist_ok=True)

logging.basicConfig(level=logging.INFO)

#######################################3
INDEX_HTML = """
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>RX PDF Tools</title>
  <link href="https://cdn.jsdelivr.net/npm/tailwindcss@2.2.19/dist/tailwind.min.css" rel="stylesheet">
</head>
<body class="bg-gray-100">
  <div class="container mx-auto p-6">
    <h1 class="text-3xl font-bold text-center text-blue-600 mb-6">RX PDF Tools</h1>
    {% if cleared_msg %}
    <div class="bg-green-100 border border-green-200 text-green-700 p-3 rounded mb-4">
      {{ cleared_msg }}
    </div>
    {% endif %}
    <div class="grid grid-cols-1 md:grid-cols-2 gap-4">
      <div class="bg-white p-4 rounded shadow">
        <h2 class="font-semibold mb-2">Smart Split & Merge (PDF Toolbox)</h2>
        <a class="inline-block mt-2 px-4 py-2 bg-blue-500 text-white rounded" href="/smart_split_merge">Open Tool</a>
      </div>
      <div class="bg-white p-4 rounded shadow">
        <h2 class="font-semibold mb-2">Saved Files</h2>
        <a class="inline-block mt-2 px-4 py-2 bg-green-600 text-white rounded" href="/saved_files">View Saved Files</a>
      </div>
      <div class="bg-white p-4 rounded shadow">
        <h2 class="font-semibold mb-2">Excel Comparator</h2>
        <a class="inline-block mt-2 px-4 py-2 bg-purple-500 text-white rounded" href="/excel">Go to Excel Tool</a>
      </div>
    </div>
  </div>
</body>
</html>
"""

#######################################
SMART_HTML = """
<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width,initial-scale=1">
  <title>Smart PDF Toolbox</title>
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
  <style>
    body { padding-bottom: 160px; } /* space for sticky bar */
    .thumb-wrapper { display:inline-block; margin:4px; }
    .sticky-actions {
      position: fixed; left: 0; right: 0; bottom: 0;
      background: #ffffff; border-top: 1px solid #e5e7eb;
      padding: 10px 16px; z-index: 9999;
      box-shadow: 0 -2px 8px rgba(0,0,0,0.05);
    }
    #pagesList { max-height: 360px; overflow:auto; }
    #thumbs { max-height: 360px; overflow:auto; white-space: normal; }
    .page-row { display:flex; justify-content:space-between; align-items:center; padding:10px 8px; border-bottom:1px solid #eee; background:#fff; border-radius:6px; margin-bottom:8px; }
    .rotate-btn { border: none; background: transparent; cursor: pointer; font-size:18px; }
    .angle-display { margin-left:8px; font-weight:600; color:#333; }
  </style>
</head>
<body class="p-4">
  <div class="container">
    <h3>Smart PDF Toolbox</h3>
    <div class="mb-3">
      <label for="pdfUpload" class="form-label">Upload PDF</label>
      <input class="form-control" type="file" id="pdfUpload" accept="application/pdf">
    </div>

    <div class="row">
      <div class="col-md-3">
        <div class="d-flex justify-content-between align-items-center mb-2">
          <h6>Pages</h6>
          <div>
            <button id="selectAllBtn" class="btn btn-sm btn-outline-primary">Select All</button>
            <button id="clearAllBtn" class="btn btn-sm btn-outline-secondary ms-1">Clear</button>
          </div>
        </div>
        <div id="pagesList" style="height:360px; overflow:auto; border:1px solid #ddd; padding:10px; background:#fdfdfd"></div>
      </div>

      <div class="col-md-6">
        <h6>Preview Thumbnails</h6>
        <div id="thumbs" style="height:360px; overflow:auto; background:#f8f9fa; padding:6px"></div>
      </div>

      <div class="col-md-3">
        <h6>Generated / Actions</h6>
        <ul id="generatedList" class="list-group mb-2"></ul>

        <div class="mb-2">
          <button id="pdfToWordBtn" class="btn btn-outline-secondary btn-sm w-100 mb-1">Export Selected → Word</button>
          <button id="pdfToExcelBtn" class="btn btn-outline-secondary btn-sm w-100">Export Selected → Excel</button>
        </div>

        <div class="small text-muted">Rotate by page numbers (e.g. <code>1,3,5-7</code>) or use per-page buttons to the right of each page.</div>
        <input id="rotatePagesInput" class="form-control form-control-sm mb-1 mt-1" placeholder="e.g. 1,3,5-7">
      </div>
    </div>
  </div>

  <div class="sticky-actions">
    <div class="container d-flex flex-column flex-md-row align-items-start gap-2">
      <div class="d-flex flex-wrap align-items-center gap-2">
        <div class="form-check form-check-inline">
          <input class="form-check-input action-checkbox" type="checkbox" value="compress" id="act_compress">
          <label class="form-check-label" for="act_compress">Compress</label>
        </div>
        <div class="form-check form-check-inline">
          <input class="form-check-input action-checkbox" type="checkbox" value="merge" id="act_merge">
          <label class="form-check-label" for="act_merge">Merge Selected Pages</label>
        </div>
        <div class="form-check form-check-inline">
          <input class="form-check-input action-checkbox" type="checkbox" value="split" id="act_split">
          <label class="form-check-label" for="act_split">Split into pages</label>
        </div>
        <div class="form-check form-check-inline">
          <input class="form-check-input action-checkbox" type="checkbox" value="remove" id="act_remove">
          <label class="form-check-label" for="act_remove">Remove Selected Pages</label>
        </div>
        <div class="form-check form-check-inline">
          <input class="form-check-input action-checkbox" type="checkbox" value="rotate" id="act_rotate">
          <label class="form-check-label" for="act_rotate">Rotate (by page numbers or per-page)</label>
        </div>

        <select id="rotateAngle" class="form-select form-select-sm ms-2" style="width:auto;">
          <option value="90">90°</option>
          <option value="180">180°</option>
          <option value="270">270°</option>
        </select>
      </div>

      <div class="d-flex flex-wrap align-items-center gap-2 ms-auto">
        <button id="generateBtn" class="btn btn-primary">Generate</button>
      </div>
    </div>
  </div>

<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/2.14.305/pdf.min.js"></script>
<script>
let uploadedFilename = '';
let selectedPages = [];
let pageRotations = {}; // per-page rotation angles (pageNum -> angle)

// Upload file to backend
async function uploadFile(file) {
  const fd = new FormData();
  fd.append('pdf', file);
  const res = await fetch('/smart_split_merge/upload', { method:'POST', body: fd });
  if (!res.ok) { alert('Upload failed'); return null }
  const data = await res.json();
  if (data && data.filename) {
    uploadedFilename = data.filename;
    loadPages(data.total_pages);
    loadThumbnails(`/smart_split_merge/uploads/${uploadedFilename}`, data.total_pages);
  } else {
    alert('Upload failed');
  }
}

// build pages list (checkboxes + per-page rotate buttons)
function loadPages(total) {
  const container = document.getElementById('pagesList');
  container.innerHTML = '';
  selectedPages = [];
  pageRotations = {};
  for (let i=1;i<=total;i++){
    const div = document.createElement('div');
    div.className = 'page-row';
    div.innerHTML = `
      <div>
        <input class="form-check-input page-checkbox me-2" type="checkbox" value="${i}" id="p${i}">
        <label class="form-check-label" for="p${i}">Page ${i}</label>
      </div>
      <div>
        <button class="btn btn-sm btn-outline-secondary rotate-btn" data-page="${i}" title="Rotate page by 90°">
          <svg xmlns="http://www.w3.org/2000/svg" width="16" height="16" fill="currentColor" class="bi bi-arrow-clockwise" viewBox="0 0 16 16">
            <path fill-rule="evenodd" d="M8 3a5 5 0 1 0 4.546 2.914.5.5 0 0 1 .908-.417A6 6 0 1 1 8 2v1z"/>
            <path d="M8 0a.5.5 0 0 1 .5.5V3a.5.5 0 0 1-1 0V.5A.5.5 0 0 1 8 0z"/>
          </svg>
        </button>
        <span id="angle-${i}" class="angle-display">0°</span>
      </div>
    `;
    container.appendChild(div);
  }

  // Checkbox event listener (delegated by class)
  document.querySelectorAll('.page-checkbox').forEach(cb=>{
    cb.addEventListener('change', function(e){
      const val = parseInt(this.value, 10);
      if (this.checked) {
        if (!selectedPages.includes(val)) selectedPages.push(val);
      } else {
        selectedPages = selectedPages.filter(x=>x!==val);
      }
    });
  });

  // rotate button handler
  document.querySelectorAll('.rotate-btn').forEach(btn=>{
    btn.addEventListener('click', function(e){
      const p = parseInt(this.getAttribute('data-page'),10);
      if (!pageRotations[p]) pageRotations[p] = 0;
      pageRotations[p] = (pageRotations[p] + 90) % 360;
      document.getElementById(`angle-${p}`).innerText = pageRotations[p] + '°';
    });
  });
}

// load thumbnails using pdf.js (client-side)
async function loadThumbnails(url, total) {
  try {
    const pdf = await pdfjsLib.getDocument(url).promise;
    const thumbs = document.getElementById('thumbs');
    thumbs.innerHTML = '';
    for (let i=1;i<=total;i++){
      const page = await pdf.getPage(i);
      const viewport = page.getViewport({ scale: 0.7 });
      const canvas = document.createElement('canvas');
      canvas.width = viewport.width;
      canvas.height = viewport.height;
      await page.render({ canvasContext: canvas.getContext('2d'), viewport }).promise;
      const wrapper = document.createElement('div');
      wrapper.className = 'thumb-wrapper';
      wrapper.appendChild(canvas);
      thumbs.appendChild(wrapper);
    }
  } catch (e) {
    console.warn('thumbnail load failed', e);
  }
}

document.getElementById('pdfUpload').addEventListener('change', e=>{
  const f = e.target.files[0];
  if (!f) return;
  uploadFile(f);
});

// Select All / Clear All
document.getElementById('selectAllBtn').addEventListener('click', ()=>{
  document.querySelectorAll('.page-checkbox').forEach(cb=>{ cb.checked = true; cb.dispatchEvent(new Event('change')); });
});
document.getElementById('clearAllBtn').addEventListener('click', ()=>{
  document.querySelectorAll('.page-checkbox').forEach(cb=>{ cb.checked = false; cb.dispatchEvent(new Event('change')); });
});

// actions getter
function getCheckedActions(){
  const actions = [];
  document.querySelectorAll('.action-checkbox:checked').forEach(n=>actions.push(n.value));
  return actions;
}

// Parse page numbers string into array of ints (1-based)
function parsePagesInput(str){
  if (!str) return [];
  const out = new Set();
  str.split(',').map(s=>s.trim()).forEach(part=>{
    if (!part) return;
    if (part.includes('-')){
      const [a,b] = part.split('-').map(x=>parseInt(x,10));
      if (!isNaN(a) && !isNaN(b)){
        for (let i=Math.min(a,b); i<=Math.max(a,b); i++) out.add(i);
      }
    } else {
      const n = parseInt(part,10);
      if (!isNaN(n)) out.add(n);
    }
  });
  return Array.from(out).sort((a,b)=>a-b);
}

// Generate button (compress/merge/split/remove/rotate)
document.getElementById('generateBtn').addEventListener('click', async ()=>{
  const actions = getCheckedActions();
  if (!uploadedFilename) { alert('Please upload a PDF first'); return }
  if (actions.length===0) { alert('Select at least one action'); return }

  // combine rotate selections: input field and per-page
  const rotatePagesInput = document.getElementById('rotatePagesInput').value;
  const rotatePagesParsed = parsePagesInput(rotatePagesInput);

  // build rotations object: prefer per-page pageRotations; also apply rotate_angle (global) to pages listed in rotatePagesParsed
  const rotations = {};
  // per-page
  for (const [p, angle] of Object.entries(pageRotations)) {
    rotations[String(p)] = angle;
  }
  // from text input (apply rotateAngle value)
  if (rotatePagesParsed && rotatePagesParsed.length > 0) {
    const globalAngle = parseInt(document.getElementById('rotateAngle').value || '90', 10);
    rotatePagesParsed.forEach(p => {
      rotations[String(p)] = ((rotations[String(p)] || 0) + globalAngle) % 360;
    });
  }

  const payload = {
    filename: uploadedFilename,
    actions: actions,
    pages: selectedPages,
    rotations: rotations
  };

  const resp = await fetch('/smart_split_merge/process', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify(payload)
  });
  const data = await resp.json();
  const list = document.getElementById('generatedList');
  list.innerHTML = '';
  (data.generated || []).forEach(fn=>{
    const li = document.createElement('li');
    li.className = 'list-group-item d-flex justify-content-between align-items-center';
    li.innerHTML = `<span>${fn}</span><div>
      <a href='/smart_split_merge/generated/${fn}' target='_blank' class='btn btn-sm btn-info me-2'>Preview</a>
      <a href='/smart_split_merge/generated/${fn}' class='btn btn-sm btn-success me-2' download>Download</a>
      </div>`;
    list.appendChild(li);
  });
});

// PDF -> Word (selected pages)####################
document.getElementById('pdfToWordBtn').addEventListener('click', async ()=>{
  if (!uploadedFilename) { alert('Upload a PDF first'); return; }
  if (selectedPages.length === 0) { alert('Select pages to export'); return; }
  const resp = await fetch('/smart_split_merge/pdf_to_word', {
    method: 'POST', headers: {'Content-Type':'application/json'},
    body: JSON.stringify({ filename: uploadedFilename, pages: selectedPages })
  });
  const data = await resp.json();
  if (data.generated && data.generated.length) {
    const gen = document.getElementById('generatedList');
    data.generated.forEach(fn=>{
      const li = document.createElement('li');
      li.className = 'list-group-item';
      li.innerHTML = `${fn} <a class="btn btn-sm btn-success float-end" href="/smart_split_merge/generated/${fn}" download>Download</a>`;
      gen.prepend(li);
    });
  } else {
    alert('Conversion failed');
  }
});

// PDF -> Excel (selected pages)####################
document.getElementById('pdfToExcelBtn').addEventListener('click', async ()=>{
  if (!uploadedFilename) { alert('Upload a PDF first'); return; }
  if (selectedPages.length === 0) { alert('Select pages to export'); return; }
  const resp = await fetch('/smart_split_merge/pdf_to_excel', {
    method: 'POST', headers: {'Content-Type':'application/json'},
    body: JSON.stringify({ filename: uploadedFilename, pages: selectedPages })
  });
  const data = await resp.json();
  if (data.generated && data.generated.length) {
    const gen = document.getElementById('generatedList');
    data.generated.forEach(fn=>{
      const li = document.createElement('li');
      li.className = 'list-group-item';
      li.innerHTML = `${fn} <a class="btn btn-sm btn-success float-end" href="/smart_split_merge/generated/${fn}" download>Download</a>`;
      gen.prepend(li);
    });
  } else {
    alert('Conversion failed');
  }
});
</script>
</body>
</html>
"""


########################3
def timestamped_name(prefix: str, ext: str = "") -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    uid = uuid.uuid4().hex[:6]
    return f"{prefix}_{ts}_{uid}{ext}"


####################################
def compress_pdf(input_path, output_path, quality="default"):
    compression_settings = {
        "default": pikepdf.CompressionLevel.default,
        "fast": pikepdf.CompressionLevel.fast,
        "none": pikepdf.CompressionLevel.none,
    }
    try:
        with pikepdf.open(input_path) as pdf:
            pdf.save(
                output_path,
                compression=compression_settings.get(
                    quality, pikepdf.CompressionLevel.default
                ),
                object_stream_mode=pikepdf.ObjectStreamMode.generate,
                linearize=True,
            )
    except Exception as e:
        raise RuntimeError(f"Compression failed: {e}")


###################################
def merge_pages(input_path: str, output_path: str, pages: list):
    reader = PdfReader(input_path)
    writer = PdfWriter()
    for p in pages:
        idx = int(p) - 1
        if 0 <= idx < len(reader.pages):
            writer.add_page(reader.pages[idx])
    with open(output_path, "wb") as f:
        writer.write(f)


#########################################
def split_pdf(input_path: str, out_folder: str) -> list:
    reader = PdfReader(input_path)
    out_files = []
    for i, page in enumerate(reader.pages, start=1):
        writer = PdfWriter()
        writer.add_page(page)
        name = timestamped_name(f"split_page_{i}", ".pdf")
        out_path = os.path.join(out_folder, name)
        with open(out_path, "wb") as f:
            writer.write(f)
        out_files.append(os.path.basename(out_path))
    return out_files


####################################33
def remove_pages(input_path: str, output_path: str, pages_to_remove: list):
    reader = PdfReader(input_path)
    writer = PdfWriter()
    remove_set = set(int(p) for p in pages_to_remove)
    for i, page in enumerate(reader.pages, start=1):
        if i not in remove_set:
            writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)


###################################
def rotate_pages_pypdf2(input_path: str, output_path: str, pages: list, angle: int):
    reader = PdfReader(input_path)
    writer = PdfWriter()
    rotate_set = set(int(p) for p in pages)
    for i, page in enumerate(reader.pages, start=1):
        if i in rotate_set:
            try:
                page.rotate_clockwise(angle)
            except Exception:
                try:
                    page.rotate(angle)
                except Exception:
                    logging.exception("rotate page fallback failed")
        writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)


####################################
def pdf_pages_to_word(input_pdf_path: str, pages: list, out_path: str):
    docx = Document()
    doc = fitz.open(input_pdf_path)
    for p in pages:
        idx = int(p) - 1
        if idx < 0 or idx >= len(doc):
            continue
        page = doc[idx]
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        img_bytes = pix.tobytes("png")
        image_stream = io.BytesIO(img_bytes)
        try:
            docx.add_picture(image_stream)
        except Exception:
            tmp = os.path.join(
                app.config["THUMBNAIL_FOLDER"], f"tmp_{uuid.uuid4().hex}.png"
            )
            with open(tmp, "wb") as f:
                f.write(img_bytes)
            docx.add_picture(tmp)
            os.remove(tmp)
        docx.add_paragraph(f"--- Page {p} ---")
        docx.add_page_break()
    docx.save(out_path)
    doc.close()


##########################################
def pdf_pages_to_excel(input_pdf_path: str, pages: list, out_path: str):
    with pdfplumber.open(input_pdf_path) as pdf:
        writer = pd.ExcelWriter(out_path, engine="openpyxl")
        for p in pages:
            idx = int(p) - 1
            if idx < 0 or idx >= len(pdf.pages):
                continue
            page = pdf.pages[idx]
            tables = page.extract_tables()
            if tables and len(tables) > 0:
                for ti, tbl in enumerate(tables):
                    if len(tbl) >= 1:
                        header = tbl[0]
                        rows = tbl[1:]
                        df = pd.DataFrame(rows, columns=header)
                    else:
                        df = pd.DataFrame(tbl)
                    sheet = f"Page_{p}_tbl{ti+1}"
                    df.to_excel(writer, sheet_name=sheet[:31], index=False)
            else:
                text = page.extract_text() or ""
                df = pd.DataFrame({"text": text.splitlines()})
                sheet = f"Page_{p}"
                df.to_excel(writer, sheet_name=sheet[:31], index=False)
        writer.close()


#####################
def parse_page_numbers_string(page_str: str) -> list:
    if not page_str:
        return []
    out = set()
    for part in page_str.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            try:
                a, b = part.split("-", 1)
                a = int(a)
                b = int(b)
                for i in range(min(a, b), max(a, b) + 1):
                    out.add(i)
            except Exception:
                continue
        else:
            try:
                out.add(int(part))
            except Exception:
                continue
    return sorted(out)


#########################
def clear_folders():
    for folder in FOLDERS_TO_CLEAR:
        try:
            if os.path.exists(folder):
                for entry in os.listdir(folder):
                    path = os.path.join(folder, entry)
                    try:
                        if os.path.isfile(path) or os.path.islink(path):
                            os.remove(path)
                        elif os.path.isdir(path):
                            shutil.rmtree(path)
                    except Exception as e:
                        logging.warning(f"Failed to delete {path}: {e}")
            else:
                os.makedirs(folder, exist_ok=True)
        except Exception as e:
            logging.exception(f"Error clearing folder {folder}: {e}")


def ensure_folders():
    for folder in FOLDERS_TO_CLEAR:
        os.makedirs(folder, exist_ok=True)


#############################3
def highlight_differences(file1, file2, merge_keys, compare_columns):
    df1 = pd.read_excel(file1)
    df2 = pd.read_excel(file2)

    merged = pd.merge(
        df1, df2, on=merge_keys, suffixes=("_file1", "_file_2"), how="inner"
    )

    red_fill = PatternFill(start_color="FF0000", end_color="FF0000", fill_type="solid")

    wb1 = load_workbook(file1)
    ws1 = wb1.active
    for _, row in merged.iterrows():
        key_values = tuple(row[k] for k in merge_keys)
        for r in ws1.iter_rows(min_row=2, values_only=False):
            row_keys = tuple(r[i].value for i in range(len(merge_keys)))
            if row_keys == key_values:
                for col in compare_columns:
                    if row[f"{col}_file1"] != row[f"{col}_file_2"]:
                        for cell in ws1[1]:
                            if cell.value == col:
                                r[cell.col_idx - 1].fill = red_fill
    file1_out = os.path.join(app.config["RESULT_FOLDER"], "comparison_file1.xlsx")
    wb1.save(file1_out)

    wb2 = load_workbook(file2)
    ws2 = wb2.active
    for _, row in merged.iterrows():
        key_values = tuple(row[k] for k in merge_keys)
        for r in ws2.iter_rows(min_row=2, values_only=False):
            row_keys = tuple(r[i].value for i in range(len(merge_keys)))
            if row_keys == key_values:
                for col in compare_columns:
                    if row[f"{col}_file1"] != row[f"{col}_file_2"]:
                        for cell in ws2[1]:
                            if cell.value == col:
                                r[cell.col_idx - 1].fill = red_fill
    file2_out = os.path.join(app.config["RESULT_FOLDER"], "comparison_file2.xlsx")
    wb2.save(file2_out)

    return file1_out, file2_out


####################################
EXCEL_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Excel Comparison Tool</title>
    <style>

        body { font-family: Arial, sans-serif; background: #f8f9fa; text-align: center; padding: 40px; }
        h2 { color: #333; margin-bottom: 20px; }
        form { background: white; padding: 20px; border-radius: 12px; max-width: 600px; margin: auto; box-shadow: 0 4px 10px rgba(0,0,0,0.1); }
        input[type="file"], button, select { width: 90%; padding: 10px; margin: 10px 0; border-radius: 8px; border: 1px solid #ccc; font-size: 14px; }
        button { background: #007bff; color: white; font-size: 16px; border: none; cursor: pointer; transition: 0.3s; }
        button:hover { background: #0056b3; }
        .section-title { font-weight: bold; color: #555; margin-top: 15px; margin-bottom: 5px; text-align: left; }
        .checkbox-group { display: flex; flex-wrap: wrap; justify-content: flex-start; gap: 10px; margin: 10px 0; }
        .checkbox-group label { background: #e9ecef; padding: 5px 10px; border-radius: 6px; cursor: pointer; font-size: 14px; }
        .checkbox-group input { margin-right: 5px; }
        .result { margin-top: 30px; background: white; padding: 20px; border-radius: 12px; max-width: 500px; margin: 30px auto; box-shadow: 0 4px 10px rgba(0,0,0,0.1); }
        a { text-decoration: none; color: #007bff; font-weight: bold; }
        a:hover { text-decoration: underline; }
        .hidden { display: none; }
        .loading { margin: 15px 0; color: #555; font-size: 14px; }
        .error { color: red; margin: 10px 0; }
    </style>
</head>
<body>
    <h2>Excel Comparison Tool</h2>
    <form id="excelForm" method="POST" enctype="multipart/form-data">
        <label>Upload File 1:</label>
        <input type="file" name="file1" id="file1" accept=".xlsx,.xls,.csv" required>

        <label>Upload File 2:</label>
        <input type="file" name="file2" id="file2" accept=".xlsx,.xls,.csv" required>

        <div id="loading" class="loading hidden">Loading common columns...</div>
        <div id="error" class="error hidden"></div>

        <div id="columns-section" class="hidden">
            <div class="section-title">Select Merge Columns:</div>
            <div class="checkbox-group" id="merge-columns"></div>

            <div class="section-title">Select Compare Columns:</div>
            <div class="checkbox-group" id="compare-columns"></div>
        </div>

        <button type="submit">Compare</button>
    </form>

    {% if download1 and download2 %}
    <div class="result">
        <h3>Download Results:</h3>
        <p><a href="/excel/download/{{ download1 }}">📥 Download File 1 Comparison</a></p>
        <p><a href="/excel/download/{{ download2 }}">📥 Download File 2 Comparison</a></p>
    </div>
    {% endif %}

    <script>
        const file1Input = document.getElementById('file1');
        const file2Input = document.getElementById('file2');
        const loadingDiv = document.getElementById('loading');
        const errorDiv = document.getElementById('error');
        const columnsSection = document.getElementById('columns-section');
        const mergeDiv = document.getElementById('merge-columns');
        const compareDiv = document.getElementById('compare-columns');

        file2Input.addEventListener('change', function () {
            errorDiv.classList.add('hidden');
            loadingDiv.classList.remove('hidden');

            let formData = new FormData();
            let file1 = file1Input.files[0];
            let file2 = file2Input.files[0];

            if (!file1 || !file2) {
                errorDiv.textContent = "Please upload both files.";
                errorDiv.classList.remove('hidden');
                loadingDiv.classList.add('hidden');
                return;
            }

            formData.append('file1', file1);
            formData.append('file2', file2);

            fetch('/excel/get_common_columns', {
                method: 'POST',
                body: formData
            })
            .then(response => response.json())
            .then(data => {
                loadingDiv.classList.add('hidden');

                if (data.error) {
                    errorDiv.textContent = data.error;
                    errorDiv.classList.remove('hidden');
                    return;
                }

                mergeDiv.innerHTML = '';
                compareDiv.innerHTML = '';

                data.columns.forEach(col => {
                    mergeDiv.innerHTML += `<label><input type="checkbox" name="merge_keys" value="${col}"> ${col}</label>`;
                    compareDiv.innerHTML += `<label><input type="checkbox" name="compare_columns" value="${col}"> ${col}</label>`;
                });

                columnsSection.classList.remove('hidden');
            })
            .catch(err => {
                loadingDiv.classList.add('hidden');
                errorDiv.textContent = "Error fetching columns: " + err;
                errorDiv.classList.remove('hidden');
            });
        });
    </script>
</body>
</html>
"""


#########################
@app.route("/")
def index():
    clear_folders()
    ensure_folders()
    cleared_msg = (
        f"Temporary files cleared at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    return render_template_string(INDEX_HTML, cleared_msg=cleared_msg)


@app.route("/smart_split_merge")
def smart_index():
    return render_template_string(SMART_HTML)


################################
@app.route("/smart_split_merge/upload", methods=["POST"])
def smart_upload():
    file = request.files.get("pdf")
    if not file:
        return jsonify({"success": False, "error": "No file uploaded"}), 400

    filename = secure_filename(file.filename)
    sub = datetime.now().strftime("%Y-%m-%d")
    upload_dir = os.path.join(app.config["UPLOAD_FOLDER"], sub)
    os.makedirs(upload_dir, exist_ok=True)
    out_path = os.path.join(upload_dir, filename)
    file.save(out_path)

    try:
        reader = PdfReader(out_path)
        total = len(reader.pages)
    except Exception:
        logging.exception("Failed to read uploaded pdf")
        return jsonify({"success": False, "error": "Invalid PDF"}), 400

    rel = os.path.join(sub, filename).replace("\\", "/")
    return jsonify({"success": True, "filename": rel, "total_pages": total})


############################
def locate_uploaded_file(rel_path):

    candidate = os.path.join(app.config["UPLOAD_FOLDER"], rel_path)
    if os.path.exists(candidate):
        return candidate
    base = os.path.basename(rel_path)
    for root, _, files in os.walk(app.config["UPLOAD_FOLDER"]):
        if base in files:
            return os.path.join(root, base)
    return None


######################################
@app.route("/smart_split_merge/process", methods=["POST"])
def smart_process():
    try:
        data = request.get_json(force=True)
        filename_rel = data.get("filename")
        actions = data.get("actions", [])
        pages_selected = data.get("pages", []) or []
        rotations = data.get("rotations", {}) or {}
        compress_flag = "compress" in actions
        do_merge = "merge" in actions
        do_split = "split" in actions
        do_remove = "remove" in actions
        do_rotate = "rotate" in actions

        if not filename_rel:
            return jsonify({"error": "missing filename"}), 400

        input_path = locate_uploaded_file(filename_rel)
        if not input_path or not os.path.exists(input_path):
            return jsonify({"error": "uploaded file not found"}), 404

        base_name = os.path.splitext(os.path.basename(filename_rel))[0]
        generated = []

        working_path = input_path
        split_done = False
        # rotation is here
        if do_rotate and rotations:
            rotated_tmp = os.path.join(
                app.config["OUTPUT_FOLDER"],
                timestamped_name(f"{base_name}_rotated", ".pdf"),
            )
            try:
                pdf_doc = fitz.open(working_path)
                for p_str, angle in rotations.items():
                    try:
                        pno = int(p_str) - 1
                        a = int(angle) % 360
                        if 0 <= pno < len(pdf_doc):
                            pdf_doc[pno].set_rotation(a)
                    except Exception:
                        logging.exception(f"Bad rotation entry {p_str}:{angle}")
                pdf_doc.save(rotated_tmp)
                pdf_doc.close()
                working_path = rotated_tmp
            except Exception:
                # if fitz fails than use pypdf2
                try:
                    reader = PdfReader(working_path)
                    writer = PdfWriter()
                    rotate_map = {int(k): int(v) % 360 for k, v in rotations.items()}
                    for i, page in enumerate(reader.pages, start=1):
                        if i in rotate_map:
                            try:
                                page.rotate_clockwise(rotate_map[i])
                            except Exception:
                                try:
                                    page.rotate(rotate_map[i])
                                except Exception:
                                    logging.exception("PyPDF2 rotate fallback failed")
                        writer.add_page(page)
                    rotated_tmp = os.path.join(
                        app.config["OUTPUT_FOLDER"],
                        timestamped_name(f"{base_name}_rotated_py2", ".pdf"),
                    )
                    with open(rotated_tmp, "wb") as f:
                        writer.write(f)
                    working_path = rotated_tmp
                except Exception:
                    logging.exception("Both rotation methods failed")

        # remove selected pages
        if do_remove and pages_selected:
            removed_tmp = os.path.join(
                app.config["OUTPUT_FOLDER"],
                timestamped_name(f"{base_name}_removed", ".pdf"),
            )
            try:
                reader = PdfReader(working_path)
                writer = PdfWriter()
                remove_set = set(int(p) for p in pages_selected)
                for i, page in enumerate(reader.pages, start=1):
                    if i not in remove_set:
                        writer.add_page(page)
                with open(removed_tmp, "wb") as f:
                    writer.write(f)
                working_path = removed_tmp
            except Exception:
                logging.exception("remove pages failed")

        # merge pages
        if do_merge and pages_selected:
            merged_tmp = os.path.join(
                app.config["OUTPUT_FOLDER"],
                timestamped_name(f"{base_name}_merged", ".pdf"),
            )
            try:
                merge_pages(working_path, merged_tmp, pages_selected)
                working_path = merged_tmp
            except Exception:
                logging.exception("merge pages failed")

        # split pages
        if do_split:
            try:
                split_files = split_pdf(working_path, app.config["OUTPUT_FOLDER"])
                for fn in split_files:
                    src = os.path.join(app.config["OUTPUT_FOLDER"], fn)
                    dst = os.path.join(app.config["SAVED_FOLDER"], fn)
                    shutil.copy(src, dst)
                    generated.append(fn)
                split_done = True
            except Exception:
                logging.exception("split failed")

        # compress with others
        if compress_flag and not split_done:
            compressed_tmp = os.path.join(
                app.config["OUTPUT_FOLDER"],
                timestamped_name(f"{base_name}_compressed", ".pdf"),
            )
            try:
                compress_pdf(working_path, compressed_tmp)
                working_path = compressed_tmp
            except Exception:
                logging.exception("compress failed")

        # final save
        try:
            final_name = timestamped_name(f"{base_name}_final", ".pdf")
            final_path = os.path.join(app.config["SAVED_FOLDER"], final_name)
            shutil.copy(working_path, final_path)
            generated.append(final_name)
        except Exception:
            logging.exception("final save failed")

        return jsonify({"generated": generated})

    except Exception as e:
        logging.exception("Processing failed")
        return jsonify({"error": str(e)}), 500


###########################3


@app.route("/smart_split_merge/pdf_to_word", methods=["POST"])
def route_pdf_to_word():
    data = request.get_json(force=True)
    filename = data.get("filename")
    pages = data.get("pages", [])
    if not filename or not pages:
        return jsonify({"error": "missing filename or pages"}), 400
    input_path = locate_uploaded_file(filename)
    if not input_path:
        return jsonify({"error": "file not found"}), 404
    out_name = timestamped_name("pdf2word", ".docx")
    out_path = os.path.join(app.config["SAVED_FOLDER"], out_name)
    try:
        pdf_pages_to_word(input_path, pages, out_path)
        return jsonify({"generated": [out_name]})
    except Exception:
        logging.exception("pdf->word failed")
        return jsonify({"error": "conversion failed"}), 500


####################################
@app.route("/smart_split_merge/pdf_to_excel", methods=["POST"])
def route_pdf_to_excel():
    data = request.get_json(force=True)
    filename = data.get("filename")
    pages = data.get("pages", [])
    if not filename or not pages:
        return jsonify({"error": "missing filename or pages"}), 400
    input_path = locate_uploaded_file(filename)
    if not input_path:
        return jsonify({"error": "file not found"}), 404
    out_name = timestamped_name("pdf2excel", ".xlsx")
    out_path = os.path.join(app.config["SAVED_FOLDER"], out_name)
    try:
        pdf_pages_to_excel(input_path, pages, out_path)
        return jsonify({"generated": [out_name]})
    except Exception:
        logging.exception("pdf->excel failed")
        return jsonify({"error": "conversion failed"}), 500


#####################################################
@app.route("/excel", methods=["GET", "POST"])
def excel_tool():
    download1, download2 = None, None
    default_merge = ["NDC"]
    default_compare = [
        "MEDID",
        "Drug Tier",
        "PA",
        "QL",
        "PA Type",
        "PA Description",
        "QL Days",
        "QL Amount",
    ]

    if request.method == "POST":
        file1 = request.files.get("file1")
        file2 = request.files.get("file2")

        if not file1 or not file2:
            return jsonify({"error": "Both files are required"}), 400

        # Save uploaded files with timestamped names
        filename1 = timestamped_name("uploaded_file1", ".xlsx")
        filename2 = timestamped_name("uploaded_file2", ".xlsx")
        path1 = os.path.join(app.config["SAVED_FOLDER"], filename1)
        path2 = os.path.join(app.config["SAVED_FOLDER"], filename2)
        file1.save(path1)
        file2.save(path2)

        # Retrieve selected merge and compare columns (from checkboxes or fallback)
        merge_keys = request.form.getlist("merge_keys") or default_merge
        compare_columns = request.form.getlist("compare_columns") or default_compare

        # Process and generate results
        result_path1, result_path2 = highlight_differences(
            path1, path2, merge_keys, compare_columns
        )

        # Save output files
        out_name1 = timestamped_name("comparison_file1", ".xlsx")
        out_name2 = timestamped_name("comparison_file2", ".xlsx")
        saved_result1 = os.path.join(app.config["RESULT_FOLDER"], out_name1)
        saved_result2 = os.path.join(app.config["RESULT_FOLDER"], out_name2)

        shutil.copy(result_path1, saved_result1)
        shutil.copy(result_path2, saved_result2)

        download1 = os.path.basename(saved_result1)
        download2 = os.path.basename(saved_result2)

    return render_template_string(
        EXCEL_TEMPLATE,
        download1=download1,
        download2=download2,
        default_merge=", ".join(default_merge),
        default_compare=", ".join(default_compare),
    )


"""@app.route('/excel/download/<filename>')
def download_file(filename):
    folder = app.config['SAVED_FOLDER']
    file_path = os.path.join(folder, filename)

    if not os.path.exists(file_path):
        return f"Error: File '{filename}' not found.", 404

    return send_from_directory(folder, filename, as_attachment=True)"""


@app.route("/excel/get_common_columns", methods=["POST"])
def get_common_columns():
    """Find and return common columns between two uploaded Excel files."""
    if "file1" not in request.files or "file2" not in request.files:
        return jsonify({"error": "Please upload both files"}), 400

    file1 = request.files["file1"]
    file2 = request.files["file2"]

    path1 = os.path.join(app.config["SAVED_FOLDER"], file1.filename)
    path2 = os.path.join(app.config["SAVED_FOLDER"], file2.filename)

    file1.save(path1)
    file2.save(path2)

    try:
        df1 = pd.read_excel(path1)
        df2 = pd.read_excel(path2)

        common_columns = sorted(list(set(df1.columns).intersection(set(df2.columns))))
        return jsonify({"columns": common_columns})
    except Exception as e:
        return jsonify({"error": f"Error reading Excel files: {str(e)}"}), 500


##################
@app.route("/smart_split_merge/uploads/<path:filename>")
def smart_uploads_serv(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)


####################3
@app.route("/smart_split_merge/generated/<path:filename>")
def smart_generated_serv(filename):
    return send_from_directory(app.config["SAVED_FOLDER"], filename)


##########################
@app.route("/saved_files")
def saved_files():
    files = os.listdir(app.config["SAVED_FOLDER"])
    lines = [
        f'<li>{f} - <a href="/smart_split_merge/generated/{f}" target="_blank">Preview</a> | <a href="/smart_split_merge/generated/{f}" download>Download</a> | <a href="/delete/{f}">Delete</a></li>'
        for f in files
    ]
    return "<h2>Saved Files</h2><ul>" + "".join(lines) + "</ul>"


##################
@app.route("/uploaded_files")
def uploaded_files():
    entries = []
    for root, _, files in os.walk(app.config["UPLOAD_FOLDER"]):
        for fn in files:
            full = os.path.join(root, fn)
            rel = os.path.relpath(full, app.config["UPLOAD_FOLDER"]).replace("\\", "/")
            entries.append(rel)
    lines = [
        f'<li><a href="/smart_split_merge/uploads/{e}">{e}</a></li>' for e in entries
    ]
    return "<h2>Uploaded Files</h2><ul>" + "".join(lines) + "</ul>"


#######################3
@app.route("/delete/<filename>")
def delete_file(filename):
    path = os.path.join(app.config["SAVED_FOLDER"], filename)
    if os.path.exists(path):
        os.remove(path)
    return redirect(url_for("saved_files"))


@app.route("/excel/download/<filename>")
def excel_download(filename):
    return send_file(
        os.path.join(app.config["RESULT_FOLDER"], filename), as_attachment=True
    )


if __name__ == "__main__":
    ensure_folders()
    app.run(debug=True, port=5000)
