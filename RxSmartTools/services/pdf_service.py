"""PDF related service helpers."""

from __future__ import annotations

import io
import logging
import shutil
import subprocess
from pathlib import Path
from typing import Iterable, Sequence

import fitz
import pandas as pd
import pdfplumber
import pikepdf
# from pikepdf import Pdf, Settings
from docx import Document
# from PyPDF2 import PdfReader, PdfWriter
from pypdf import PdfReader, PdfWriter

from ..utils.filesystem import timestamped_name

LOGGER = logging.getLogger(__name__)


def compress_pdf(input_path: Path, output_path: Path, quality: str = "default") -> None:
    """Compress ``input_path`` and write the result to ``output_path``."""
    # If user explicitly requests no compression, just copy the file.
    print(f'Quality : {quality}')
    if quality == "none":
        shutil.copy(input_path, output_path)
        return

     

    #  rewrite using PyPDF2 (may not reduce size significantly) or copy
    try: 
        writer = PdfWriter(clone_from=input_path)

        # for page in writer.pages:
        #     page.compress_content_streams()

        for page in writer.pages:
            for img in page.images:
                img.replace(img.image, quality=50)

        with open(output_path, "wb") as handle:
            writer.write(handle)

        print(f"output_path : {output_path}")

    except Exception:
        print("PyPDF2 compression failed, trying pikepdf")
        LOGGER.exception("Final fallback failed; copying original file")
        shutil.copy(input_path, output_path)


def merge_pages(
    input_path: Path, page_numbers: Sequence[int], output_path: Path
) -> None:
    """Create a PDF containing ``page_numbers`` extracted from ``input_path``."""

    reader = PdfReader(input_path)
    writer = PdfWriter()
    for number in page_numbers:
        index = number - 1
        if 0 <= index < len(reader.pages):
            writer.add_page(reader.pages[index])
    with output_path.open("wb") as handle:
        writer.write(handle)


def merge_files(
    input_paths: Sequence[Path], output_path: Path, progress_callback=None, cancel_check=None
) -> None:
    """Merge whole input files into a single PDF stored at ``output_path``.

    - input_paths: sequence of Paths to source PDFs (order preserved)
    - progress_callback: optional callable(msg: str) for reporting progress
    - cancel_check: optional callable() -> bool; if returns True, raise RuntimeError to cancel
    """
    writer = PdfWriter()
    first_meta = None
    for idx, p in enumerate(input_paths, start=1):
        if cancel_check and cancel_check():
            raise RuntimeError("Merge cancelled")
        if progress_callback:
            progress_callback(f"Merging file {idx}/{len(input_paths)}: {p.name}")
        reader = PdfReader(p)
        # preserve metadata from first document when available
        if first_meta is None:
            try:
                first_meta = reader.metadata
            except Exception:
                first_meta = None
        for page in reader.pages:
            writer.add_page(page)

    if first_meta:
        try:
            writer.add_metadata(first_meta)
        except Exception:
            LOGGER.exception("Failed to add metadata to merged output")

    with output_path.open("wb") as handle:
        writer.write(handle)

    # quick sanity-check of output
    try:
        out_reader = PdfReader(output_path)
        if len(out_reader.pages) == 0:
            raise RuntimeError("Merged output has no pages")
    except Exception:
        LOGGER.exception("Validation of merged output failed")
        raise


def merge_selected_pages_from_multiple_files(
    file_page_map: dict[Path, Sequence[int]],
    output_path: Path,
    progress_callback=None,
    cancel_check=None,
) -> None:
    """Create a single merged PDF from selected pages across multiple files.

    file_page_map: mapping of Path -> iterable of 1-based page numbers
    """
    writer = PdfWriter()
    first_meta = None
    total_files = len(file_page_map)
    for idx, (p, pages) in enumerate(file_page_map.items(), start=1):
        if cancel_check and cancel_check():
            raise RuntimeError("Merge cancelled")
        if progress_callback:
            progress_callback(f"Processing {idx}/{total_files}: {p.name}")
        reader = PdfReader(p)
        if first_meta is None:
            try:
                first_meta = reader.metadata
            except Exception:
                first_meta = None

        # validate pages
        for number in pages:
            index = int(number) - 1
            if not (0 <= index < len(reader.pages)):
                raise ValueError(f"Invalid page {number} for file {p.name}")
            writer.add_page(reader.pages[index])

    if first_meta:
        try:
            writer.add_metadata(first_meta)
        except Exception:
            LOGGER.exception("Failed to add metadata to merged output")

    with output_path.open("wb") as handle:
        writer.write(handle)

    # validate
    try:
        out_reader = PdfReader(output_path)
        if len(out_reader.pages) == 0:
            raise RuntimeError("Merged output has no pages")
    except Exception:
        LOGGER.exception("Validation of merged output failed")
        raise


def split_pdf(input_path: Path, destination_dir: Path) -> list[Path]:
    """Split ``input_path`` into individual pages stored under ``destination_dir``."""

    reader = PdfReader(input_path)
    created_files: list[Path] = []
    for index, page in enumerate(reader.pages, start=1):
        writer = PdfWriter()
        writer.add_page(page)
        file_name = timestamped_name(f"split_page_{index}", ".pdf")
        output_path = destination_dir / file_name
        with output_path.open("wb") as handle:
            writer.write(handle)
        created_files.append(output_path)
    return created_files


def remove_pages(
    input_path: Path, pages_to_remove: Iterable[int], output_path: Path
) -> None:
    """Remove the supplied ``pages_to_remove`` from ``input_path``."""

    reader = PdfReader(input_path)
    writer = PdfWriter()
    removal_set = {int(page_number) for page_number in pages_to_remove}
    for index, page in enumerate(reader.pages, start=1):
        if index not in removal_set:
            writer.add_page(page)
    with output_path.open("wb") as handle:
        writer.write(handle)


def rotate_pages(
    input_path: Path,
    rotations: dict[int, int],
    output_path: Path,
) -> None:
    """Rotate individual pages according to ``rotations`` mapping."""

    try:
        rotate_with_fitz(input_path, rotations, output_path)
    except Exception:  # pragma: no cover - fallback path
        LOGGER.exception("Primary rotation failed; falling back to PyPDF2")
        rotate_with_pypdf2(input_path, rotations, output_path)


def rotate_with_fitz(
    input_path: Path,
    rotations: dict[int, int],
    output_path: Path,
) -> None:
    """Rotate pages using PyMuPDF."""

    document = fitz.open(input_path)
    try:
        for page_number, angle in rotations.items():
            index = page_number - 1
            if 0 <= index < len(document):
                document[index].set_rotation(angle % 360)
        document.save(output_path)
    finally:
        document.close()


def rotate_with_pypdf2(
    input_path: Path,
    rotations: dict[int, int],
    output_path: Path,
) -> None:
    """Rotate pages using PyPDF2 as a compatibility fallback."""

    reader = PdfReader(input_path)
    writer = PdfWriter()
    for index, page in enumerate(reader.pages, start=1):
        angle = rotations.get(index)
        if angle is not None:
            try:
                page.rotate_clockwise(angle)
            except Exception:  # pragma: no cover
                LOGGER.exception("PyPDF2 rotate_clockwise failed; trying rotate")
                page.rotate(angle)
        writer.add_page(page)
    with output_path.open("wb") as handle:
        writer.write(handle)


def pages_to_word(
    input_pdf: Path, page_numbers: Sequence[int], output_path: Path
) -> None:
    """Create a DOCX document containing rasterized ``page_numbers`` from ``input_pdf``."""

    document = Document()
    pdf = fitz.open(input_pdf)
    try:
        for page_number in page_numbers:
            index = page_number - 1
            if not (0 <= index < len(pdf)):
                continue
            page = pdf[index]
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_bytes = pixmap.tobytes("png")
            buffer = io.BytesIO(image_bytes)
            try:
                document.add_picture(buffer)
            except Exception:  # pragma: no cover - temporary file fallback
                LOGGER.exception("Inline DOCX picture failed; using temp file")
                temp_path = output_path.parent / timestamped_name("page", ".png")
                with temp_path.open("wb") as temp_file:
                    temp_file.write(image_bytes)
                document.add_picture(str(temp_path))
                temp_path.unlink(missing_ok=True)
            document.add_paragraph(f"--- Page {page_number} ---")
            document.add_page_break()
        document.save(output_path)
    finally:
        pdf.close()


def pages_to_excel(
    input_pdf: Path, page_numbers: Sequence[int], output_path: Path
) -> None:
    """Export table data from ``page_numbers`` to an Excel workbook."""

    with pdfplumber.open(input_pdf) as pdf:
        with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
            for page_number in page_numbers:
                index = page_number - 1
                if not (0 <= index < len(pdf.pages)):
                    continue
                page = pdf.pages[index]
                tables = page.extract_tables() or []
                if tables:
                    for table_index, table in enumerate(tables, start=1):
                        dataframe = _table_to_dataframe(table)
                        sheet_name = f"Page_{page_number}_tbl{table_index}"[:31]
                        dataframe.to_excel(writer, sheet_name=sheet_name, index=False)
                else:
                    text = page.extract_text() or ""
                    dataframe = pd.DataFrame({"text": text.splitlines()})
                    sheet_name = f"Page_{page_number}"[:31]
                    dataframe.to_excel(writer, sheet_name=sheet_name, index=False)


def copy_outputs_to_saved(files: Iterable[Path], saved_dir: Path) -> list[str]:
    """Copy generated files into ``saved_dir`` and return their filenames."""

    saved_dir.mkdir(parents=True, exist_ok=True)
    generated: list[str] = []
    for source in files:
        destination = saved_dir / source.name
        shutil.copy(source, destination)
        generated.append(destination.name)
    return generated


def _table_to_dataframe(table: list[list[str]]) -> pd.DataFrame:
    """Convert ``table`` data extracted by pdfplumber into a DataFrame."""

    if not table:
        return pd.DataFrame()
    header, *rows = table
    if header:
        return pd.DataFrame(rows, columns=header)
    return pd.DataFrame(table)
